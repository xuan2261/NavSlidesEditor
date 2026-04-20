import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    # Setup paths
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    docs_dir = os.path.join(base_dir, 'docs', 'manual')
    md_path = os.path.join(docs_dir, 'content', 'detailed_manual.md')
    screenshots_dir = os.path.join(docs_dir, 'screenshots')
    output_path = os.path.join(docs_dir, 'NavSlides_User_Manual.docx')

    # Load content
    if not os.path.exists(md_path):
        print(f"Error: Could not find {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Initialize document
    doc = Document()

    # Set base style to Times New Roman, 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Create Cover Page
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title_text = 'HƯỚNG DẪN SỬ DỤNG PHẦN MỀM NAVSLIDES'
    title = doc.add_paragraph(title_text)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    author = doc.add_paragraph('Tác giả: Bùi Thanh Xuân\nKhoa Kỹ thuật cơ sở - Học viện Hải quân')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

    version = doc.add_paragraph("Phiên bản: 1.0 (2026)")
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in version.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()

    # Helper function to add image safely
    def add_screenshot(filename):
        img_path = os.path.join(screenshots_dir, filename)
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(img_path, width=Inches(6.0))
            
            caption = doc.add_paragraph(f"Hình: {filename.replace('.png', '').replace('_', ' ').title()}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.style = doc.styles['Caption']
        else:
            print(f"Warning: Image {filename} not found.")

    # Parse Markdown
    lines = md_content.split('\n')
    current_chapter = ""
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Main Title, skip because cover page is already done
            if "HƯỚNG DẪN SỬ DỤNG" in line.upper():
                continue
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            # Chapter
            if current_chapter != "":
                doc.add_page_break()
            current_chapter = line.replace('## ', '')
            doc.add_heading(current_chapter, level=1)
        elif line.startswith('### '):
            # Section
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('[SCREENSHOT:'):
            # Extract filename
            filename = line.replace('[SCREENSHOT:', '').replace(']', '').strip()
            add_screenshot(filename)
        elif line.startswith('> '):
            # Blockquote
            p = doc.add_paragraph(line.replace('> ', ''))
            p.style = 'Intense Quote'
        elif line.startswith('- '):
            # Bullet point
            doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        else:
            # Normal text
            doc.add_paragraph(line)

    # Save Document
    doc.save(output_path)
    print(f"[Success] Generated rich manual at: {output_path}")

if __name__ == '__main__':
    create_manual()
